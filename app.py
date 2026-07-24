import streamlit as st
import os
import time
from datetime import datetime
from io import BytesIO
from dotenv import load_dotenv

load_dotenv()

st.set_page_config(
    page_title="AI Marketing Automation Agent",
    page_icon="🚀",
    layout="wide"
)

if 'theme' not in st.session_state:
    st.session_state.theme = 'dark'

def set_theme():
    if st.session_state.theme == 'dark':
        st.markdown("""
        <style>
        .stApp { background: #0e1117; }
        .stSidebar { background: #1e1e2e; }
        h1, h2, h3, .stTitle, .stHeader, .stSubheader { color: #ffffff !important; }
        .stTextInput > label, .stSelectbox > label, .stTextArea > label { color: #ffffff !important; }
        .stTextInput > div > div > input {
            background: #2d2d44 !important;
            color: #ffffff !important;
            border: 1px solid #444 !important;
            border-radius: 8px !important;
            padding: 10px 14px !important;
        }
        .stTextInput > div > div > input::placeholder { color: #888 !important; }
        .stSelectbox > div > div > select {
            background: #2d2d44 !important;
            color: #ffffff !important;
            border: 1px solid #444 !important;
            border-radius: 8px !important;
            padding: 10px 14px !important;
        }
        .stTextArea > div > div > textarea {
            background: #2d2d44 !important;
            color: #ffffff !important;
            border: 1px solid #444 !important;
            border-radius: 8px !important;
        }
        .stTextArea > div > div > textarea::placeholder { color: #888 !important; }
        .stButton > button {
            background: #ff4b4b !important;
            color: white !important;
            border: none !important;
            border-radius: 8px !important;
            padding: 10px 24px !important;
            font-weight: 600 !important;
        }
        .stButton > button:hover { background: #ff6b6b !important; }
        .stDownloadButton > button {
            background: #2d6a4f !important;
            color: white !important;
            border: none !important;
            border-radius: 8px !important;
            padding: 8px 16px !important;
        }
        .stDownloadButton > button:hover { background: #40916c !important; }
        .stTabs [data-baseweb="tab-list"] {
            background: #1e1e2e !important;
            border-radius: 8px !important;
            padding: 4px !important;
        }
        .stTabs [data-baseweb="tab"] {
            color: #8899aa !important;
            border-radius: 6px !important;
            padding: 8px 20px !important;
        }
        .stTabs [aria-selected="true"] {
            background: #ff4b4b !important;
            color: #ffffff !important;
        }
        .stWarning { background: #332700 !important; color: #ffd700 !important; }
        .stError { background: #3b0000 !important; color: #ff6b6b !important; }
        .stSuccess { background: #003300 !important; color: #00ff00 !important; }
        .stMarkdown { color: #e0e0e0 !important; }
        hr { border-color: #444 !important; }
        .stCaption { color: #888 !important; }
        .stMetric { background: #1e1e2e !important; border-radius: 8px !important; padding: 10px !important; }
        .stMetric label { color: #ffffff !important; }
        </style>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <style>
        .stApp { background: #ffffff; }
        .stSidebar { background: #f0f2f6; }
        h1, h2, h3 { color: #000000 !important; }
        .stTextInput > label { color: #000000 !important; }
        .stTextInput > div > div > input { background: #ffffff !important; color: #000000 !important; border: 1px solid #ccc !important; border-radius: 8px !important; padding: 10px 14px !important; }
        .stButton > button { background: #ff4b4b !important; color: white !important; border: none !important; border-radius: 8px !important; padding: 10px 24px !important; }
        .stTabs [aria-selected="true"] { background: #ff4b4b !important; color: #ffffff !important; }
        </style>
        """, unsafe_allow_html=True)

set_theme()

try:
    from docx import Document
    DOCX_AVAILABLE = True
except:
    DOCX_AVAILABLE = False

# ========== GROQ API - DIRECT HARDCODED KEY (FINAL) ==========
groq_available = False
client = None

# ✅ DIRECT KEY - NO SECRETS, NO ENV
GROQ_API_KEY = "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"

try:
    from groq import Groq
    client = Groq(api_key=GROQ_API_KEY)
    groq_available = True
    print("✅ Groq Connected Successfully!")
except Exception as e:
    print(f"❌ Groq Error: {e}")
    groq_available = False

def generate_with_groq(prompt):
    if not groq_available:
        return "❌ Groq API not available. Check your API keys."
    
    models_to_try = [
       'llama3-70b-8192',
    'llama3-8b-8192',
    'mixtral-8x7b-32768',
    'gemma2-9b-it',
    ]
    
    for model_name in models_to_try:
        try:
            response = client.chat.completions.create(
                model=model_name,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
                max_tokens=2048
            )
            return response.choices[0].message.content
        except Exception as e:
            continue
    
    return "❌ No models available. Please try again."

with st.sidebar:
    st.markdown("## ⚙️ Configuration")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🌙 Dark", use_container_width=True):
            st.session_state.theme = 'dark'
            st.rerun()
    with col2:
        if st.button("☀️ Light", use_container_width=True):
            st.session_state.theme = 'light'
            st.rerun()
    
    st.divider()
    st.markdown("### 🏢 Business Details")
    
    business_name = st.text_input("Business Name", placeholder="e.g., TechSolutions", key="biz_input")
    industry = st.text_input("Industry", placeholder="e.g., Software Development", key="ind_input")
    target_audience = st.text_input("Target Audience", placeholder="e.g., Small Business Owners", key="aud_input")
    country = st.text_input("Country", placeholder="e.g., India", key="cnt_input")
    brand_name = st.text_input("Brand Name", placeholder="e.g., TechSol", key="brd_input")
    tone = st.selectbox("Tone of Content", ["Professional", "Casual", "Funny", "Inspirational", "Formal"], key="tone_input")
    
    st.divider()
    st.markdown("### 🔑 API Status")
    if groq_available:
        st.success("✅ Groq API Connected")
    else:
        st.error("❌ No API Available")

st.markdown("# 🚀 AI Marketing Automation Agent")
st.markdown("Generate marketing content, hashtags, emails, and campaigns using AI.")

def is_valid():
    if not business_name or not industry:
        st.warning("⚠️ Please fill Business Name and Industry")
        return False
    if not groq_available:
        st.error("❌ No API available. Check your API keys.")
        return False
    return True

def generate(prompt_type, extra=""):
    biz = business_name or "My Business"
    ind = industry or "Industry"
    aud = target_audience or "general audience"
    cnt = country or "global"
    brd = brand_name or biz
    tn = tone or "Professional"
    
    prompts = {
        "content": f"Generate marketing content for {biz} ({ind}). Audience: {aud}. Country: {cnt}. Brand: {brd}. Tone: {tn}. Include social media post, email, blog intro, CTA.",
        "ideas": f"Generate 5 marketing campaign ideas for {biz} ({ind}). Audience: {aud}. Country: {cnt}. Brand: {brd}. Tone: {tn}.",
        "hashtags": f"Generate 30 hashtags for {biz} ({ind}). Audience: {aud}. Country: {cnt}. Brand: {brd}.",
        "email": f"Write email campaign for {biz} ({ind}). Audience: {aud}. Country: {cnt}. Brand: {brd}. Tone: {tn}.",
        "custom": f"{extra}\n\nBusiness: {biz} ({ind}). Audience: {aud}. Country: {cnt}. Brand: {brd}. Tone: {tn}."
    }
    
    with st.spinner("🔄 Generating..."):
        time.sleep(0.3)
        return generate_with_groq(prompts.get(prompt_type, prompts["custom"]))

def make_docx(content, title):
    if not DOCX_AVAILABLE:
        return None
    try:
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph("")
        
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except:
        return None

def show_output(content, title):
    st.markdown(content)
    
    safe_title = "".join(c for c in title if c.isalnum() or c in " _-")[:50]
    filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M')}"
    
    docx_data = make_docx(content, title)
    if docx_data:
        st.download_button(
            label="📝 Download DOCX",
            data=docx_data,
            file_name=f"{filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    else:
        st.warning("⚠️ DOCX unavailable")

t1, t2, t3, t4, t5 = st.tabs(["📝 Content", "💡 Ideas", "#️⃣ Hashtags", "✉️ Email", "🎨 Custom"])

with t1:
    st.subheader("📝 Generate Marketing Content")
    if is_valid():
        if st.button("🚀 Generate Content", use_container_width=True):
            result = generate("content")
            show_output(result, f"Content_{business_name}")

with t2:
    st.subheader("💡 Marketing Campaign Ideas")
    if is_valid():
        if st.button("🚀 Generate Ideas", use_container_width=True):
            result = generate("ideas")
            show_output(result, f"Ideas_{business_name}")

with t3:
    st.subheader("#️⃣ Generate Hashtags")
    if is_valid():
        if st.button("🚀 Generate Hashtags", use_container_width=True):
            result = generate("hashtags")
            show_output(result, f"Hashtags_{business_name}")

with t4:
    st.subheader("✉️ Email Campaign")
    if is_valid():
        if st.button("🚀 Generate Email", use_container_width=True):
            result = generate("email")
            show_output(result, f"Email_{business_name}")

with t5:
    st.subheader("🎨 Custom Prompt")
    if is_valid():
        custom = st.text_area("Your prompt", height=120, placeholder="Write a LinkedIn post...")
        if st.button("🚀 Generate", use_container_width=True) and custom:
            result = generate("custom", custom)
            show_output(result, f"Custom_{business_name}")

st.divider()
st.caption("🚀 AI Marketing Automation Agent | Powered by Groq API")

with st.expander("ℹ️ System Status"):
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.metric("DOCX", "✅" if DOCX_AVAILABLE else "❌")
    with c2:
        st.metric("Groq API", "✅" if groq_available else "❌")
    with c3:
        st.metric("Theme", "🌙" if st.session_state.theme == 'dark' else "☀️")
    with c4:
        st.metric("PDF", "❌ Removed")
